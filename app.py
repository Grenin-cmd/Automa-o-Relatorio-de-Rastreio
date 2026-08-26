from __future__ import annotations

import io
from io import BytesIO
import math
import os
import socket
import sqlite3
import sys
import tempfile
import threading
import time
import unicodedata
import uuid
import webbrowser
from typing import Any

from dotenv import load_dotenv

# 1. Carrega o .env antes de qualquer import de agente
if getattr(sys, "frozen", False):
    BASE_DIR = os.path.dirname(os.path.abspath(sys.executable))
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

env_path = os.path.join(BASE_DIR, ".env")
load_dotenv(dotenv_path=env_path, override=True)

# 2. Imports das bibliotecas e Flask
import numpy as np
import pandas as pd
from flask import Flask, flash, redirect, render_template, request, send_file, session, url_for

try:
    from docx import Document  # type: ignore[reportMissingImports]
except ImportError:  # pragma: no cover
    Document = None

# 3. Módulos do sistema (agora seguros, com a chave do Gemini já carregada)
from analisador_rastreio import RastreamentoAnalyzer
from agente_refinamento import agente_ia


def _template_folder() -> str:
    if getattr(sys, "frozen", False):
        return os.path.join(sys._MEIPASS, "templates")  # type: ignore[attr-defined]
    return "templates"


app = Flask(__name__, template_folder=_template_folder())
app.secret_key = "mudar_para_uma_chave_secreta"


def _persistent_data_dir() -> str:
    if getattr(sys, "frozen", False):
        return os.path.dirname(os.path.abspath(sys.executable))
    return os.path.dirname(os.path.abspath(__file__))

def _get_db_connection():
    db_path = os.path.join(_persistent_data_dir(), "poi_store.db")
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn


def _init_poi_db() -> None:
    conn = _get_db_connection()
    try:
        colunas_existentes = {row[1]: row[2] for row in conn.execute("PRAGMA table_info(pois)").fetchall()}
        if colunas_existentes and colunas_existentes.get("id", "").upper() != "TEXT":
            conn.execute("DROP TABLE IF EXISTS pois")
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS pois (
                id TEXT PRIMARY KEY,
                nome TEXT NOT NULL,
                tipo TEXT,
                latitude REAL NOT NULL,
                longitude REAL NOT NULL,
                cidade TEXT,
                estado TEXT,
                rodovia TEXT,
                raio_metros REAL,
                tempo_parado_seg INTEGER,
                velocidade_maxima_kmh REAL
            )
            """
        )
        conn.execute("CREATE INDEX IF NOT EXISTS idx_pois_nome ON pois(nome COLLATE NOCASE);")
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS regras_refinamento (
                id TEXT PRIMARY KEY,
                poi_id TEXT,
                categoria TEXT NOT NULL,
                descricao TEXT NOT NULL,
                ativa INTEGER DEFAULT 1,
                criada_em TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (poi_id) REFERENCES pois(id) ON DELETE CASCADE
            )
            """
        )
        conn.execute("CREATE INDEX IF NOT EXISTS idx_regras_ativa ON regras_refinamento(ativa);")
        conn.commit()
    finally:
        conn.close()


_init_poi_db()


def _get_regras_ativas() -> list[dict[str, Any]]:
    conn = _get_db_connection()
    try:
        rows = conn.execute(
            "SELECT * FROM regras_refinamento WHERE ativa = 1 ORDER BY criada_em DESC"
        ).fetchall()
        return [dict(row) for row in rows]
    finally:
        conn.close()


@app.context_processor
def inject_globais():
    conn = _get_db_connection()
    pois = []
    try:
        cur = conn.execute("SELECT id, nome, raio_metros, latitude, longitude FROM pois ORDER BY nome ASC")
        pois = [dict(row) for row in cur.fetchall()]
    except Exception:
        pass
    finally:
        conn.close()
    return dict(regras_salvas=_get_regras_ativas(), pois_salvos=pois)


def _normalize_column_name(name: str) -> str:
    name = str(name).strip().lower()
    name = unicodedata.normalize("NFKD", name)
    name = "".join(ch for ch in name if not unicodedata.combining(ch))
    name = name.replace(" ", "_").replace("/", "_").replace("-", "_")
    while "__" in name:
        name = name.replace("__", "_")
    return name


def _normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [_normalize_column_name(col) for col in df.columns]
    return df


def _parse_number(value: Any, cast: type, error_message: str):
    if pd.isna(value):
        raise ValueError(error_message)

    if isinstance(value, str):
        normalized = value.strip().replace("\xa0", " ").replace("\u202f", " ").replace(" ", "")
        if not normalized or normalized.lower() in {"nan", "n/a", "na", "none", "null", "-", "--"}:
            raise ValueError(error_message)

        if "," in normalized:
            if "." in normalized and normalized.rfind(".") < normalized.rfind(","):
                normalized = normalized.replace(".", "")
            normalized = normalized.replace(",", ".")
        elif normalized.count(".") > 1:
            first_dot = normalized.find(".")
            normalized = normalized[:first_dot+1] + normalized[first_dot+1:].replace(".", "")

        value = normalized

    try:
        return cast(value)
    except Exception:
        raise ValueError(error_message)


def _resolver_coluna(df: pd.DataFrame, aliases: list[str]) -> str:
    columns = [_normalize_column_name(col) for col in df.columns]
    for alias in aliases:
        alias_norm = _normalize_column_name(alias)
        for i, col in enumerate(columns):
            if not col:
                continue
            if alias_norm == col or alias_norm in col or col in alias_norm:
                return df.columns[i]
    raise ValueError(f"Coluna não encontrada. Esperava uma das opções: {aliases}")


def ler_pois(texto: str) -> list[dict[str, object]]:
    pois = []
    for linha in texto.splitlines():
        if not linha.strip():
            continue
        partes = linha.strip().split(":")
        if len(partes) != 7:
            continue
        nome, tipo, lat, lon, raio, tempo, velocidade = partes
        pois.append(
            {
                "nome": nome,
                "tipo": tipo,
                "latitude": _parse_number(lat, float, "Latitude ou longitude inválida no arquivo de POIs."),
                "longitude": _parse_number(lon, float, "Latitude ou longitude inválida no arquivo de POIs."),
                "raio_metros": _parse_number(raio, float, "Raio inválido no arquivo de POIs."),
                "tempo_parado_seg": _parse_number(tempo, int, "Tempo parado inválido no arquivo de POIs."),
                "velocidade_maxima_kmh": _parse_number(velocidade, float, "Velocidade máxima inválida no arquivo de POIs."),
            }
        )
    return pois


def _save_uploaded_file(uploaded_file) -> str:
    dados = uploaded_file.read()
    _, extensao = os.path.splitext(uploaded_file.filename)
    extensao = extensao.lower() if extensao else ".csv"
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=extensao)
    tmp.write(dados)
    tmp.flush()
    tmp.close()
    return tmp.name


def _adicionar_regra(descricao: str, categoria: str = "GERAL", poi_id: str | None = None) -> str:
    regra_id = uuid.uuid4().hex
    conn = _get_db_connection()
    try:
        conn.execute(
            """
            INSERT INTO regras_refinamento (id, poi_id, categoria, descricao, ativa)
            VALUES (?, ?, ?, ?, 1)
            """,
            (regra_id, poi_id, categoria, descricao),
        )
        conn.commit()
        return regra_id
    finally:
        conn.close()


@app.route("/excluir_regras", methods=["POST"])
def excluir_regras():
    regra_ids = request.form.getlist("regra_ids")
    if not regra_ids:
        flash("Nenhuma regra selecionada para exclusão.", "warning")
        return redirect("/")

    conn = _get_db_connection()
    try:
        placeholders = ",".join("?" for _ in regra_ids)
        conn.execute(f"DELETE FROM regras_refinamento WHERE id IN ({placeholders})", tuple(regra_ids))
        conn.commit()
        flash(f"{len(regra_ids)} regra(s) excluída(s) com sucesso!", "success")
    except Exception as e:
        flash(f"Erro ao excluir regras: {e}", "danger")
    finally:
        conn.close()
    return redirect("/")


@app.route("/excluir_pois", methods=["POST"])
def excluir_pois():
    poi_ids = request.form.getlist("poi_ids")
    if not poi_ids:
        flash("Nenhum POI selecionado para exclusão.", "warning")
        return redirect("/")

    conn = _get_db_connection()
    try:
        placeholders = ",".join("?" for _ in poi_ids)
        conn.execute(f"DELETE FROM pois WHERE id IN ({placeholders})", tuple(poi_ids))
        conn.commit()
        flash(f"{len(poi_ids)} POI(s) excluído(s) com sucesso!", "success")
    except Exception as e:
        flash(f"Erro ao excluir POIs: {e}", "danger")
    finally:
        conn.close()
    return redirect("/")


def _create_poi_id() -> str:
    return uuid.uuid4().hex


def _get_saved_pois() -> list[dict[str, object]]:
    conn = _get_db_connection()
    try:
        rows = conn.execute("SELECT * FROM pois ORDER BY nome COLLATE NOCASE").fetchall()
        return [dict(row) for row in rows]
    finally:
        conn.close()


def _save_saved_pois(pois: list[dict[str, object]]) -> None:
    conn = _get_db_connection()
    try:
        conn.execute("DELETE FROM pois")
        conn.executemany(
            """
            INSERT INTO pois (id, nome, tipo, latitude, longitude, cidade, estado, rodovia, raio_metros, tempo_parado_seg, velocidade_maxima_kmh)
            VALUES (:id, :nome, :tipo, :latitude, :longitude, :cidade, :estado, :rodovia, :raio_metros, :tempo_parado_seg, :velocidade_maxima_kmh)
            """,
            pois,
        )
        conn.commit()
    finally:
        conn.close()


def _parse_poi_dataframe(df: pd.DataFrame) -> list[dict[str, object]]:
    if df.empty:
        raise ValueError("Arquivo de POIs vazio.")

    working = _normalize_columns(df.copy())

    def _find_column(aliases: list[str]) -> str | None:
        for alias in aliases:
            try:
                return _resolver_coluna(working, [alias])
            except ValueError:
                continue
        return None

    columns = {
        "nome": _find_column(["nome", "name", "poi", "ponto", "nome_poi", "poi_nome", "local", "localizacao", "localidade", "location"]),
        "tipo": _find_column(["tipo", "type", "categoria", "category", "grupo", "group"]),
        "latitude": _find_column(["latitude", "lat", "latidude"]),
        "longitude": _find_column(["longitude", "lon", "lng", "longitutde", "longitud", "longitute"]),
        "raio_metros": _find_column(["raio_metros", "raio", "radius", "distancia", "distance"]),
        "tempo_parado_seg": _find_column(["tempo_parado_seg", "tempo_parado", "tempo", "tempo_segundos", "stop_time"]),
        "velocidade_maxima_kmh": _find_column(["velocidade_maxima_kmh", "velocidade_maxima", "velocidade", "velocidade_kmh", "speed_kmh", "speed"]),
        "cidade": _find_column(["cidade", "city"]),
        "estado": _find_column(["estado", "state"]),
        "rodovia": _find_column(["rodovia", "highway", "road"]),
    }

    missing = [key for key in ["nome", "tipo", "latitude", "longitude"] if columns[key] is None]
    if missing:
        raise ValueError("Arquivo de POIs inválido. Faltam colunas obrigatórias: " + ", ".join(missing))

    nome_col = columns["nome"]
    latitude_col = columns["latitude"]
    longitude_col = columns["longitude"]
    tipo_col = columns["tipo"]
    cidade_col = columns["cidade"]
    estado_col = columns["estado"]
    rodovia_col = columns["rodovia"]
    raio_col = columns["raio_metros"]
    tempo_col = columns["tempo_parado_seg"]
    velocidade_col = columns["velocidade_maxima_kmh"]

    assert nome_col is not None and latitude_col is not None and longitude_col is not None

    pois: list[dict[str, object]] = []
    for _, row in working.iterrows():
        nome = str(row[nome_col]).strip()
        if not nome:
            continue

        try:
            latitude = _parse_number(row[latitude_col], float, "Latitude ou longitude inválida no arquivo de POIs.")
            longitude = _parse_number(row[longitude_col], float, "Latitude ou longitude inválida no arquivo de POIs.")
        except ValueError:
            continue

        tipo = str(row[tipo_col]).strip() if tipo_col else ""
        cidade = str(row[cidade_col]).strip() if cidade_col else ""
        estado = str(row[estado_col]).strip() if estado_col else ""
        rodovia = str(row[rodovia_col]).strip() if rodovia_col else ""

        try:
            raio_metros = _parse_number(row[raio_col], float, "") if raio_col and pd.notna(row[raio_col]) else 200.0
        except ValueError:
            raio_metros = 200.0

        try:
            tempo_parado_seg = _parse_number(row[tempo_col], int, "") if tempo_col and pd.notna(row[tempo_col]) else 120
        except ValueError:
            tempo_parado_seg = 120

        try:
            velocidade_maxima_kmh = _parse_number(row[velocidade_col], float, "") if velocidade_col and pd.notna(row[velocidade_col]) else 10.0
        except ValueError:
            velocidade_maxima_kmh = 10.0

        pois.append(
            {
                "id": _create_poi_id(),
                "nome": nome,
                "tipo": tipo,
                "latitude": latitude,
                "longitude": longitude,
                "cidade": cidade,
                "estado": estado,
                "rodovia": rodovia,
                "raio_metros": raio_metros,
                "tempo_parado_seg": tempo_parado_seg,
                "velocidade_maxima_kmh": velocidade_maxima_kmh,
            }
        )

    if not pois:
        raise ValueError("Nenhum POI válido encontrado no arquivo.")

    return pois


def _cleanup_cached_file() -> None:
    file_path = session.pop("uploaded_file_path", None)
    session.pop("uploaded_file_name", None)
    if file_path and os.path.exists(file_path):
        try:
            os.unlink(file_path)
        except OSError:
            pass


def _format_duration(seconds: int) -> str:
    minutes = seconds // 60
    if minutes < 60:
        return f"{minutes} min"
    hours = minutes // 60
    minutes = minutes % 60
    return f"{hours}h {minutes}m" if minutes else f"{hours}h"


def _parse_velocidade_series(series: pd.Series) -> pd.Series:
    extraida = series.astype(str).str.extract(r"(-?\d+(?:[.,]\d+)?)")[0]
    return pd.to_numeric(extraida.str.replace(",", ".", regex=False), errors="coerce")


def _haversine_metros(lat1: float, lon1: float, lat2: float, lon2: float) -> float:
    r = 6371000.0
    phi1, phi2 = math.radians(lat1), math.radians(lat2)
    delta_phi = math.radians(lat2 - lat1)
    delta_lambda = math.radians(lon2 - lon1)
    a = math.sin(delta_phi / 2.0) ** 2 + math.cos(phi1) * math.cos(phi2) * math.sin(delta_lambda / 2.0) ** 2
    return r * (2.0 * math.atan2(math.sqrt(a), math.sqrt(1.0 - a)))


def _find_missing_poi_rows(
    df: pd.DataFrame, 
    raio_tolerancia_m: float = 200.0, 
    nomes_pois_registrados: set[str] | None = None,
    pois_lista: list[dict[str, object]] | None = None
) -> tuple[list[dict[str, object]], list[str]]:
    try:
        timestamp_col = _resolver_coluna(df, ["timestamp", "data_hora", "datetime", "date_time", "horario", "data"])
        speed_col = _resolver_coluna(df, ["velocidade_kmh", "speed_kmh", "velocidade", "speed"])
    except ValueError:
        return [], []

    lat_col = None
    lon_col = None
    for alias in ["latitude", "lat"]:
        try:
            lat_col = _resolver_coluna(df, [alias])
            break
        except ValueError:
            continue

    for alias in ["longitude", "lon", "lng"]:
        try:
            lon_col = _resolver_coluna(df, [alias])
            break
        except ValueError:
            continue

    working = _normalize_columns(df.copy())
    working[timestamp_col] = pd.to_datetime(working[timestamp_col], dayfirst=True, format="mixed", errors="coerce")
    working[speed_col] = _parse_velocidade_series(working[speed_col]).fillna(9999)

    parados = working[working[speed_col] <= 10].dropna(subset=[timestamp_col]).copy()
    if parados.empty:
        return [], []

    if pois_lista and lat_col and lon_col:
        lats_p = pd.to_numeric(parados[lat_col].astype(str).str.replace(",", "."), errors="coerce").values
        lons_p = pd.to_numeric(parados[lon_col].astype(str).str.replace(",", "."), errors="coerce").values

        poi_coords = []
        poi_raios = []
        for p in pois_lista:
            try:
                plat = float(str(p.get("latitude")).replace(",", "."))
                plon = float(str(p.get("longitude")).replace(",", "."))
                praio = float(str(p.get("raio_metros", 200.0)).replace(",", "."))
                poi_coords.append((plat, plon))
                poi_raios.append(max(praio, float(raio_tolerancia_m)))
            except (ValueError, TypeError):
                continue

        if poi_coords:
            poi_arr = np.radians(np.array(poi_coords))
            raios_arr = np.array(poi_raios)

            rad_lats = np.radians(lats_p)[:, None]
            rad_lons = np.radians(lons_p)[:, None]
            poi_lats = poi_arr[:, 0][None, :]
            poi_lons = poi_arr[:, 1][None, :]

            dlat = poi_lats - rad_lats
            dlon = poi_lons - rad_lons
            a = np.sin(dlat / 2.0)**2 + np.cos(rad_lats) * np.cos(poi_lats) * np.sin(dlon / 2.0)**2
            c = 2.0 * np.arctan2(np.sqrt(a), np.sqrt(1.0 - a))
            dist_metros = 6371000.0 * c

            is_longe = ~np.any(dist_metros <= raios_arr[None, :], axis=1)
            missing = parados[is_longe].copy()
        else:
            missing = parados.copy()
    else:
        poi_col = None
        for alias in ["poi", "local", "location", "ponto_de_interesse"]:
            try:
                poi_col = _resolver_coluna(df, [alias])
                break
            except ValueError:
                continue

        if poi_col:
            sem_poi = parados[poi_col].isna() | parados[poi_col].astype(str).str.strip().eq("")
            missing = parados[sem_poi].copy()
        else:
            missing = parados.copy()

    missing = missing.sort_values(timestamp_col).reset_index(drop=True)
    if missing.empty:
        return [], []

    groups: list[dict[str, object]] = []
    current_start = None
    current_end = None
    current_count = 0

    for _, row in missing.iterrows():
        timestamp = row[timestamp_col]
        if current_start is None:
            current_start = timestamp
            current_end = timestamp
            current_count = 1
            continue

        if timestamp - current_end <= pd.Timedelta(minutes=10):
            current_end = timestamp
            current_count += 1
        else:
            if (current_end - current_start).total_seconds() >= 600:
                groups.append(
                    {
                        "Início": current_start,
                        "Fim": current_end,
                        "Duração": _format_duration(int((current_end - current_start).total_seconds())),
                        "Registros": current_count,
                    }
                )
            current_start = timestamp
            current_end = timestamp
            current_count = 1

    if current_start is not None and current_end is not None and (current_end - current_start).total_seconds() >= 600:
        groups.append(
            {
                "Início": current_start,
                "Fim": current_end,
                "Duração": _format_duration(int((current_end - current_start).total_seconds())),
                "Registros": current_count,
            }
        )

    if not groups:
        return [], []

    columns = ["Início", "Fim", "Duração", "Registros"]
    rows = [{str(key): value for key, value in record.items()} for record in groups]
    return rows, columns


def _find_suspicious_stops(
    df: pd.DataFrame, 
    raio_tolerancia_m: float = 200.0, 
    nomes_pois_registrados: set[str] | None = None,
    pois_lista: list[dict[str, object]] | None = None
) -> list[dict[str, object]]:
    try:
        placa_col = _resolver_coluna(df, ["Placa", "placa", "plate"])
    except ValueError:
        return []

    resultados: list[dict[str, object]] = []
    placas = sorted(df[placa_col].dropna().astype(str).unique().tolist())
    for placa in placas:
        subset = df[df[placa_col].astype(str) == placa]
        rows, _ = _find_missing_poi_rows(subset, raio_tolerancia_m, nomes_pois_registrados, pois_lista)
        for row in rows:
            resultados.append(
                {
                    "Placa": placa,
                    "Início": row["Início"],
                    "Fim": row["Fim"],
                    "Duração": row["Duração"],
                    "Registros": row["Registros"],
                }
            )

    resultados.sort(key=lambda item: str(item.get("Início", "")))
    return resultados


def _gerar_timeline_viagem(df: pd.DataFrame, placa: str, pois: list[dict[str, object]], limite_velocidade_kmh: float, raio_tolerancia_m: float = 200.0) -> list[dict[str, object]]:
    try:
        placa_col = _resolver_coluna(df, ["Placa", "placa", "plate"])
    except ValueError:
        return []

    subset = df[df[placa_col].astype(str) == placa]
    if subset.empty:
        return []

    analyzer = RastreamentoAnalyzer(pois=pois)
    eventos = list(analyzer.gerar_eventos_viagem(subset, limite_velocidade_kmh))

    nomes_pois_registrados = {_normalize_column_name(str(p["nome"])) for p in pois if p.get("nome")}
    missing_rows, _ = _find_missing_poi_rows(subset, raio_tolerancia_m, nomes_pois_registrados)
    for row in missing_rows:
        eventos.append(
            {
                "tipo": "parada_suspeita",
                "local": None,
                "inicio": row["Início"],
                "fim": row["Fim"],
                "em_andamento": False,
            }
        )

    eventos.sort(key=lambda e: e["inicio"])
    return eventos


def _gerar_resumo_geral_placas(df: pd.DataFrame, pois: list[dict[str, object]], raio_tolerancia_m: float = 200.0) -> str:
    try:
        working = _normalize_columns(df.copy())
        placa_col = _resolver_coluna(working, ["Placa", "placa", "plate"])
    except ValueError:
        return "Não foi possível detectar a coluna de placa para gerar o resumo geral."

    placas = sorted(working[placa_col].dropna().astype(str).unique().tolist())
    if not placas:
        return "Nenhuma placa encontrada no arquivo."

    analyzer = RastreamentoAnalyzer(pois=pois, raio_tolerancia_m=raio_tolerancia_m)
    lines: list[str] = []
    for placa in placas:
        subset = df[df[placa_col].astype(str) == placa]
        if subset.empty:
            continue

        report = analyzer.gerar_relatorio(subset)
        report_body = report
        if report.startswith("Relatório de rastreamento"):
            report_body = report.split("\n\n", 1)[-1]

        report_body = report_body.strip()
        if not report_body:
            report_body = "Nenhum evento relevante detectado."

        lines.append(f"{placa}: {report_body}")

    resumo_bruto = "\n\n".join(lines)
    regras_ativas = _get_regras_ativas()
    return agente_ia.refinar_relatorio(resumo_bruto, regras_ativas)


def _create_word_bytes(summary: str) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx não está instalado. Instale com 'pip install python-docx'.")

    document = Document()
    document.add_heading("Resumo Geral de Placas", level=1)
    document.add_paragraph("Análise Inteligente de Frotas", style="Intense Quote")
    for line in summary.split("\n\n"):
        document.add_paragraph(line)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.read()


@app.route("/download_word")
def download_word():
    resumo = session.get("word_summary")
    if not resumo:
        flash("Gere o resumo geral primeiro para poder baixar o Word.", "warning")
        return redirect("/")

    try:
        doc_bytes = _create_word_bytes(resumo)
    except RuntimeError as error:
        flash(str(error), "danger")
        return redirect("/")

    return send_file(
        BytesIO(doc_bytes),
        as_attachment=True,
        download_name="resumo_geral_frotas.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def carregar_arquivo(uploaded_file) -> pd.DataFrame:
    if isinstance(uploaded_file, str):
        caminho = uploaded_file
        if not os.path.exists(caminho):
            raise FileNotFoundError(f"Arquivo não encontrado: {caminho}.")

        extensao = os.path.splitext(caminho)[1].lower()
        if extensao == ".csv":
            for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin1"):
                for sep in [",", ";", "\t", "|"]:
                    try:
                        df = pd.read_csv(caminho, encoding=encoding, sep=sep)
                        df = _normalize_columns(df)
                        if len(df.columns) > 1:
                            return df
                    except Exception:
                        continue
            raise ValueError("Não foi possível decodificar o CSV. Verifique o encoding e o delimitador.")

        if extensao in {".xls", ".xlsx"}:
            try:
                df = pd.read_excel(caminho, engine="calamine")
            except Exception:
                df = pd.read_excel(caminho)
            return _normalize_columns(df)

        raise ValueError("Formato não suportado. Use CSV ou Excel.")

    nome = uploaded_file.filename
    if not nome:
        raise ValueError("Nenhum arquivo selecionado.")

    _, extensao = os.path.splitext(nome)
    extensao = extensao.lower()

    if extensao == ".csv":
        dados = uploaded_file.read()
        for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin1"):
            for sep in [",", ";", "\t", "|"]:
                try:
                    df = pd.read_csv(io.BytesIO(dados), encoding=encoding, sep=sep)
                    df = _normalize_columns(df)
                    if len(df.columns) > 1:
                        return df
                except Exception:
                    continue
        raise ValueError("Não foi possível decodificar o CSV.")

    if extensao in {".xls", ".xlsx"}:
        dados = uploaded_file.read()
        buffer = io.BytesIO(dados)
        try:
            df = pd.read_excel(buffer, engine="calamine")
        except Exception:
            df = pd.read_excel(buffer)
        return _normalize_columns(df)

    raise ValueError("Formato não suportado. Use CSV ou Excel.")


@app.route("/", methods=["GET", "POST"])
def index():
    report = None
    pois_text = (
        "ClienteA:cliente:-23.5505:-46.6333:300:120:10\n"
        "PostoX:posto:-23.5600:-46.6400:250:120:10"
    )
    placa_escolhida = request.form.get("placa", "")
    tipo_veiculo = request.form.get("tipo_veiculo", "pesado")
    try:
        raio_suspeita = float(request.form.get("raio_suspeita", 200) or 200)
    except (TypeError, ValueError):
        raio_suspeita = 200.0

    action = request.form.get("action", "analyze")
    columns: list[str] = []
    missing_rows: list[dict[str, object]] = []
    missing_columns: list[str] = []
    placas: list[str] = []
    resumo_geral = None
    paradas_suspeitas: list[dict[str, object]] = []
    timeline_eventos: list[dict[str, object]] = []
    docx_available = Document is not None
    uploaded_file_path = session.get("uploaded_file_path")
    file_name = session.get("uploaded_file_name", "")
    saved_pois = _get_saved_pois()
    saved_pois_count = len(saved_pois)
    nomes_pois_registrados = {_normalize_column_name(str(p["nome"])) for p in saved_pois if p.get("nome")}

    if uploaded_file_path and not os.path.exists(uploaded_file_path):
        _cleanup_cached_file()
        uploaded_file_path = None
        file_name = ""

    # GET RÁPIDO
    if request.method == "GET" and uploaded_file_path and os.path.exists(uploaded_file_path):
        try:
            df_full = carregar_arquivo(uploaded_file_path)
            columns = list(df_full.columns)
            placa_col = _resolver_coluna(df_full, ["Placa", "placa", "plate"])
            placas = sorted(df_full[placa_col].dropna().astype(str).unique().tolist())
            if not placa_escolhida and placas:
                placa_escolhida = placas[0]
            
            df_placa = df_full[df_full[placa_col].astype(str) == placa_escolhida] if placa_escolhida else df_full
            missing_rows, missing_columns = _find_missing_poi_rows(df_placa, raio_suspeita, nomes_pois_registrados)
            paradas_suspeitas = _find_suspicious_stops(df_placa, raio_suspeita, nomes_pois_registrados)
        except Exception:
            _cleanup_cached_file()
            uploaded_file_path = None
            file_name = ""

    # POST OTIMIZADO (CARREGAMENTO ÚNICO DO ARQUIVO)
    if request.method == "POST":
        arquivo = request.files.get("arquivo")
        arquivo_pois = request.files.get("arquivo_pois")
        pois_text = request.form.get("pois", "")
        saved_pois = _get_saved_pois()
        saved_pois_count = len(saved_pois)

        action = request.form.get("action") or request.form.get("submit") or ""

        if action == "import_pois" or "import_pois" in request.form or (arquivo_pois and arquivo_pois.filename):
            if not arquivo_pois or not arquivo_pois.filename:
                flash("Selecione um arquivo CSV ou Excel de POIs para importar.", "warning")
            else:
                try:
                    df_pois = carregar_arquivo(arquivo_pois)
                    imported_pois = _parse_poi_dataframe(df_pois)
                    _save_saved_pois(imported_pois)
                    flash(f"{len(imported_pois)} POI(s) importado(s) com sucesso.", "success")
                except Exception as error:
                    flash(f"Erro ao processar POIs: {error}", "danger")
            return redirect("/")

        if arquivo and arquivo.filename:
            try:
                new_path = _save_uploaded_file(arquivo)
                if uploaded_file_path and uploaded_file_path != new_path and os.path.exists(uploaded_file_path):
                    try:
                        os.unlink(uploaded_file_path)
                    except OSError:
                        pass
                uploaded_file_path = new_path
                session["uploaded_file_path"] = new_path
                session["uploaded_file_name"] = arquivo.filename
                file_name = arquivo.filename
            except Exception as error:
                flash(str(error), "danger")

        if not uploaded_file_path or not os.path.exists(uploaded_file_path):
            flash("Selecione um arquivo CSV ou Excel para continuar.", "warning")
            return render_template(
                "index.html",
                report=report,
                placa_escolhida=placa_escolhida,
                tipo_veiculo=tipo_veiculo,
                raio_suspeita=raio_suspeita,
                paradas_suspeitas=paradas_suspeitas,
                timeline_eventos=timeline_eventos,
                placas=placas,
                file_name=file_name,
                resumo_geral=resumo_geral,
                docx_available=docx_available,
            )

        # 1. Carrega o arquivo UMA ÚNICA VEZ em memória
        try:
            df_full = carregar_arquivo(uploaded_file_path)
            columns = list(df_full.columns)
            placa_col = _resolver_coluna(df_full, ["Placa", "placa", "plate"])
            placas = sorted(df_full[placa_col].dropna().astype(str).unique().tolist())
            if not placa_escolhida and placas:
                placa_escolhida = placas[0]
        except Exception as error:
            flash(str(error), "danger")
            return render_template(
                "index.html",
                report=report,
                placa_escolhida=placa_escolhida,
                tipo_veiculo=tipo_veiculo,
                raio_suspeita=raio_suspeita,
                paradas_suspeitas=paradas_suspeitas,
                timeline_eventos=timeline_eventos,
                placas=placas,
                file_name=file_name,
                resumo_geral=resumo_geral,
                docx_available=docx_available,
            )

        df_placa = df_full[df_full[placa_col].astype(str) == placa_escolhida] if placa_escolhida else df_full

        # 2. Processa os POIs
        pois = ler_pois(pois_text) + [{k: v for k, v in p.items() if k != "id"} for p in saved_pois]
        pois_ajustados = []
        for p in pois:
            p_copia = dict(p)
            val_raio = p_copia.get("raio_metros")
            try:
                raio_orig = float(str(val_raio).replace(",", ".")) if val_raio is not None else 200.0
            except (ValueError, TypeError):
                raio_orig = 200.0
            p_copia["raio_metros"] = max(raio_orig, float(raio_suspeita))
            pois_ajustados.append(p_copia)

        analyzer = RastreamentoAnalyzer(pois=pois_ajustados, raio_tolerancia_m=raio_suspeita)
        limite_velocidade_kmh = 95.0 if tipo_veiculo == "pesado" else 110.0
        regras_ativas = _get_regras_ativas()

        # 3. Executa a Ação Solicitada
        try:
            if action == "summary":
                resumo_geral = _gerar_resumo_geral_placas(df_full, pois_ajustados, raio_tolerancia_m=raio_suspeita)
                session["word_summary"] = resumo_geral
            elif action == "timeline":
                timeline_eventos = _gerar_timeline_viagem(df_full, placa_escolhida, pois_ajustados, limite_velocidade_kmh, raio_suspeita)
            else:
                relatorio_bruto = analyzer.gerar_relatorio(df_placa)
                report = agente_ia.refinar_relatorio(relatorio_bruto, regras_ativas)

            paradas_suspeitas = _find_suspicious_stops(df_placa, raio_suspeita, nomes_pois_registrados)
        except Exception as error:
            flash(str(error), "danger")

        return render_template(
            "index.html",
            report=report,
            placa_escolhida=placa_escolhida,
            tipo_veiculo=tipo_veiculo,
            raio_suspeita=raio_suspeita,
            paradas_suspeitas=paradas_suspeitas,
            timeline_eventos=timeline_eventos,
            placas=placas,
            file_name=file_name,
            resumo_geral=resumo_geral,
            docx_available=docx_available,
        )

    return render_template(
        "index.html",
        report=report,
        placa_escolhida="",
        tipo_veiculo=tipo_veiculo,
        raio_suspeita=raio_suspeita,
        paradas_suspeitas=paradas_suspeitas,
        timeline_eventos=timeline_eventos,
        placas=[],
        file_name=file_name,
        resumo_geral=resumo_geral,
        docx_available=docx_available,
    )


@app.route("/adicionar_regra_feedback", methods=["POST"])
def adicionar_regra_feedback():
    descricao = request.form.get("descricao_regra", "").strip()
    categoria = request.form.get("categoria", "GERAL").strip()
    poi_id = request.form.get("poi_id") or None

    if not descricao:
        flash("A descrição da regra não pode estar vazia.", "warning")
    else:
        _adicionar_regra(descricao, categoria, poi_id)
        flash("Instrução de IA salva com sucesso!", "success")

    return redirect("/")


def _find_available_port(initial_port: int = 5000) -> int:
    port = initial_port
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
        while True:
            try:
                sock.bind(("0.0.0.0", port))
                return port
            except OSError:
                port += 1


def _is_port_available(port: int) -> bool:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
        try:
            sock.bind(("0.0.0.0", port))
            return True
        except OSError:
            return False


def _get_server_port(default_port: int = 5000) -> int:
    port_str = os.environ.get("PORT")
    if port_str:
        try:
            port = int(port_str)
        except ValueError:
            raise RuntimeError("Valor inválido para a variável de ambiente PORT. Use um número inteiro.")

        if _is_port_available(port):
            return port
        raise RuntimeError(f"Porta {port} já está em uso. Pare o serviço existente ou defina PORT com outra porta.")

    if _is_port_available(default_port):
        return default_port

    return _find_available_port(default_port + 1)


def _open_browser(port: int) -> None:
    url = f"http://127.0.0.1:{port}/"

    def _try_open() -> None:
        time.sleep(1.0)
        try:
            webbrowser.open_new_tab(url)
        except Exception:
            pass

    thread = threading.Thread(target=_try_open, daemon=True)
    thread.start()


if __name__ == "__main__":
    port = _get_server_port(5000)
    print(f"Iniciando a aplicação na porta {port}...")
    _open_browser(port)

    app.run(
        host="0.0.0.0",
        port=port,
        debug=False,
        use_reloader=False,
    )